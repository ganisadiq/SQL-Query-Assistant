from src.loaders.base_loader import BaseLoader
from src.models.document import Document
from pathlib import Path
from docx import Document as DocxDocument

class DOCXLoader(BaseLoader):

    def load(self, source) -> list[Document]:
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"{source} doesnt exist")
        
        page_text = []
        
        doc = DocxDocument(path)

        for paragraph in doc.paragraphs:
            text = self._extract_paragraph_text(paragraph)
            page_text.append(text)

        full_text = "\n".join(page_text)

        documents = Document(
        document_id=path.stem,
        filename=path.name,
        text=full_text,
        metadata={
            "source":str(path),
            "type":"docx"
        }
        )

        return [documents]
    
    def _extract_paragraph_text(self, paragraph):
        return paragraph.text


